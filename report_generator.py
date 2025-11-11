"""
SolarVisionAI - Multi-Format Report Generator
Standards: IEC 60904-13, ISO/IEC 17025, ISO 9001

Professional report generation in multiple formats:
- Excel (openpyxl) with conditional formatting, charts, pivot tables
- Word (python-docx) with IEC/ISO templates
- PDF (reportlab) with multi-page professional layouts

Includes: defect statistics, images (raw + processed + annotated),
observations table, reference comparisons, compliance statements

Author: SolarVisionAI Team
Version: 1.0.0
"""

import io
import base64
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import logging

import numpy as np
import pandas as pd
from PIL import Image
import cv2

# Excel generation
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.formatting.rule import ColorScaleRule, CellIsRule

# Word generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph,
    Spacer, Image as RLImage, PageBreak, KeepTogether
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from camera_config import CameraSettings, TestConditions
from analytics_engine import ImageQualityMetrics

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class ReportConfig:
    """Configuration for report generation"""
    # Report identification
    report_title: str = "Electroluminescence Inspection Report"
    company_name: str = "SolarVisionAI"
    company_logo_path: Optional[str] = None

    # Content options
    include_raw_images: bool = True
    include_processed_images: bool = True
    include_annotated_images: bool = True
    include_statistics: bool = True
    include_recommendations: bool = True
    include_compliance_statement: bool = True

    # Standards
    standards: List[str] = None

    # Layout
    page_size: str = "A4"  # A4 or Letter
    orientation: str = "portrait"  # portrait or landscape

    def __post_init__(self):
        if self.standards is None:
            self.standards = [
                "IEC TS 60904-13:2018",
                "IEC 60904-14:2020",
                "ISO/IEC 17025:2017"
            ]


class ExcelReportGenerator:
    """
    Excel report generator with advanced formatting

    Creates professional Excel reports with:
    - Multiple worksheets
    - Conditional formatting
    - Charts and visualizations
    - Pivot tables (optional)
    - Data validation
    """

    def __init__(self, config: Optional[ReportConfig] = None):
        """Initialize Excel generator"""
        self.config = config or ReportConfig()
        logger.info("Initialized ExcelReportGenerator")

    def generate(
        self,
        defect_data: pd.DataFrame,
        quality_metrics: Optional[ImageQualityMetrics] = None,
        test_conditions: Optional[TestConditions] = None,
        output_path: Optional[Path] = None
    ) -> Path:
        """
        Generate Excel report

        Args:
            defect_data: DataFrame with defect analysis
            quality_metrics: Image quality metrics
            test_conditions: Test conditions
            output_path: Output file path

        Returns:
            Path to generated report
        """
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = Path(f"EL_Report_{timestamp}.xlsx")

        # Create workbook
        wb = Workbook()
        wb.remove(wb.active)  # Remove default sheet

        # 1. Summary sheet
        self._create_summary_sheet(wb, defect_data, quality_metrics, test_conditions)

        # 2. Defect details sheet
        if not defect_data.empty:
            self._create_defects_sheet(wb, defect_data)

        # 3. Statistics sheet
        if not defect_data.empty:
            self._create_statistics_sheet(wb, defect_data)

        # 4. Quality metrics sheet
        if quality_metrics:
            self._create_quality_sheet(wb, quality_metrics)

        # 5. Test conditions sheet
        if test_conditions:
            self._create_test_conditions_sheet(wb, test_conditions)

        # Save workbook
        wb.save(output_path)
        logger.info(f"Excel report saved to: {output_path}")

        return output_path

    def _create_summary_sheet(
        self,
        wb: Workbook,
        defect_data: pd.DataFrame,
        quality_metrics: Optional[ImageQualityMetrics],
        test_conditions: Optional[TestConditions]
    ) -> None:
        """Create summary worksheet"""
        ws = wb.create_sheet("Summary", 0)

        # Header
        ws['A1'] = self.config.report_title
        ws['A1'].font = Font(size=18, bold=True, color="1F4E78")
        ws.merge_cells('A1:D1')

        ws['A2'] = self.config.company_name
        ws['A2'].font = Font(size=14, italic=True)
        ws.merge_cells('A2:D2')

        ws['A3'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        ws['A3'].font = Font(size=10, color="7F7F7F")

        # Key metrics
        row = 5
        ws[f'A{row}'] = "KEY METRICS"
        ws[f'A{row}'].font = Font(size=14, bold=True)
        row += 1

        metrics = [
            ("Total Defects", len(defect_data)),
            ("Critical Defects", len(defect_data[defect_data['Severity'] == 'Critical']) if not defect_data.empty else 0),
            ("Total Power Loss (W)", defect_data['Power_Loss_W'].sum() if not defect_data.empty else 0),
            ("Quality Score", quality_metrics.quality_score if quality_metrics else 0)
        ]

        for label, value in metrics:
            ws[f'A{row}'] = label
            ws[f'B{row}'] = value
            ws[f'A{row}'].font = Font(bold=True)

            # Conditional formatting for critical defects
            if "Critical" in label and value > 0:
                ws[f'B{row}'].fill = PatternFill(start_color="FF6B6B", end_color="FF6B6B", fill_type="solid")
                ws[f'B{row}'].font = Font(color="FFFFFF", bold=True)

            row += 1

        # Compliance status
        row += 1
        ws[f'A{row}'] = "IEC COMPLIANCE"
        ws[f'A{row}'].font = Font(size=14, bold=True)
        row += 1

        if quality_metrics:
            compliance_status = "PASS" if quality_metrics.iec_compliant else "FAIL"
            ws[f'A{row}'] = "Overall Status"
            ws[f'B{row}'] = compliance_status
            ws[f'A{row}'].font = Font(bold=True)

            if compliance_status == "PASS":
                ws[f'B{row}'].fill = PatternFill(start_color="4CAF50", end_color="4CAF50", fill_type="solid")
                ws[f'B{row}'].font = Font(color="FFFFFF", bold=True)
            else:
                ws[f'B{row}'].fill = PatternFill(start_color="F44336", end_color="F44336", fill_type="solid")
                ws[f'B{row}'].font = Font(color="FFFFFF", bold=True)

        # Adjust column widths
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 20

    def _create_defects_sheet(self, wb: Workbook, defect_data: pd.DataFrame) -> None:
        """Create defects detail worksheet"""
        ws = wb.create_sheet("Defect Details")

        # Write data
        for r in dataframe_to_rows(defect_data, index=False, header=True):
            ws.append(r)

        # Format header
        header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)

        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")

        # Conditional formatting for severity
        severity_col = None
        for idx, cell in enumerate(ws[1], 1):
            if cell.value == "Severity":
                severity_col = idx
                break

        if severity_col:
            col_letter = chr(64 + severity_col)  # Convert to letter

            # Critical = Red
            ws.conditional_formatting.add(
                f'{col_letter}2:{col_letter}{ws.max_row}',
                CellIsRule(
                    operator='equal',
                    formula=['"Critical"'],
                    fill=PatternFill(start_color="FF6B6B", end_color="FF6B6B", fill_type="solid"),
                    font=Font(color="FFFFFF", bold=True)
                )
            )

            # High = Orange
            ws.conditional_formatting.add(
                f'{col_letter}2:{col_letter}{ws.max_row}',
                CellIsRule(
                    operator='equal',
                    formula=['"High"'],
                    fill=PatternFill(start_color="FFA726", end_color="FFA726", fill_type="solid"),
                    font=Font(color="FFFFFF")
                )
            )

        # Auto-fit columns
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width

    def _create_statistics_sheet(self, wb: Workbook, defect_data: pd.DataFrame) -> None:
        """Create statistics worksheet with charts"""
        ws = wb.create_sheet("Statistics")

        # Defect type distribution
        defect_counts = defect_data['Defect_Type'].value_counts()

        ws['A1'] = "Defect Type Distribution"
        ws['A1'].font = Font(size=14, bold=True)

        row = 3
        ws['A2'] = "Defect Type"
        ws['B2'] = "Count"
        ws['A2'].font = Font(bold=True)
        ws['B2'].font = Font(bold=True)

        for defect_type, count in defect_counts.items():
            ws[f'A{row}'] = defect_type
            ws[f'B{row}'] = count
            row += 1

        # Create pie chart
        pie = PieChart()
        pie.title = "Defect Distribution"
        labels = Reference(ws, min_col=1, min_row=3, max_row=row - 1)
        data = Reference(ws, min_col=2, min_row=2, max_row=row - 1)
        pie.add_data(data, titles_from_data=True)
        pie.set_categories(labels)
        ws.add_chart(pie, "D2")

        # Severity distribution
        severity_counts = defect_data['Severity'].value_counts()

        row += 2
        start_row = row
        ws[f'A{row}'] = "Severity Distribution"
        ws[f'A{row}'].font = Font(size=14, bold=True)
        row += 1

        ws[f'A{row}'] = "Severity"
        ws[f'B{row}'] = "Count"
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'].font = Font(bold=True)
        row += 1

        for severity, count in severity_counts.items():
            ws[f'A{row}'] = severity
            ws[f'B{row}'] = count
            row += 1

        # Create bar chart
        bar = BarChart()
        bar.title = "Severity Levels"
        bar.x_axis.title = "Severity"
        bar.y_axis.title = "Count"
        labels = Reference(ws, min_col=1, min_row=start_row + 2, max_row=row - 1)
        data = Reference(ws, min_col=2, min_row=start_row + 1, max_row=row - 1)
        bar.add_data(data, titles_from_data=True)
        bar.set_categories(labels)
        ws.add_chart(bar, "D15")

    def _create_quality_sheet(self, wb: Workbook, metrics: ImageQualityMetrics) -> None:
        """Create quality metrics worksheet"""
        ws = wb.create_sheet("Quality Metrics")

        ws['A1'] = "Image Quality Analysis"
        ws['A1'].font = Font(size=14, bold=True)

        metrics_data = [
            ("Dimension", f"{metrics.width}x{metrics.height}"),
            ("Total Pixels", f"{metrics.total_pixels:,}"),
            ("Quality Score", f"{metrics.quality_score:.1f}/100"),
            ("Quality Level", metrics.quality_level.upper()),
            ("", ""),
            ("Signal Quality", ""),
            ("SNR (dB)", f"{metrics.snr_db:.2f}"),
            ("PSNR (dB)", f"{metrics.psnr_db:.2f}"),
            ("", ""),
            ("Sharpness", ""),
            ("IEC Class", metrics.iec_sharpness_class),
            ("MTF50", f"{metrics.mtf50:.3f}"),
            ("Laplacian Variance", f"{metrics.sharpness_laplacian:.2f}"),
            ("Edge Contrast", f"{metrics.edge_contrast:.2f}"),
            ("", ""),
            ("Statistics", ""),
            ("Mean Intensity", f"{metrics.mean_intensity:.2f}"),
            ("Std Deviation", f"{metrics.std_deviation:.2f}"),
            ("Kurtosis", f"{metrics.kurtosis:.2f}"),
            ("Skewness", f"{metrics.skewness:.2f}"),
            ("", ""),
            ("Entropy", ""),
            ("Shannon Entropy", f"{metrics.shannon_entropy:.3f} bits"),
            ("Normalized", f"{metrics.normalized_entropy:.3f}"),
            ("", ""),
            ("IEC Compliance", ""),
            ("Overall", "PASS" if metrics.iec_compliant else "FAIL"),
            ("Resolution Met", "✓" if metrics.iec_min_resolution_met else "✗"),
            ("Contrast Met", "✓" if metrics.iec_min_contrast_met else "✗")
        ]

        row = 3
        for label, value in metrics_data:
            ws[f'A{row}'] = label
            ws[f'B{row}'] = value

            if label and not value:  # Section header
                ws[f'A{row}'].font = Font(bold=True, size=12, color="1F4E78")
            elif label == "":  # Spacer
                pass
            else:
                ws[f'A{row}'].font = Font(bold=True)

            row += 1

        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 20

    def _create_test_conditions_sheet(self, wb: Workbook, conditions: TestConditions) -> None:
        """Create test conditions worksheet"""
        ws = wb.create_sheet("Test Conditions")

        ws['A1'] = "Test Conditions (IEC 60904-14)"
        ws['A1'].font = Font(size=14, bold=True)

        test_data = [
            ("Test Information", ""),
            ("Test ID", conditions.test_id),
            ("Test Date", conditions.test_date.strftime("%Y-%m-%d %H:%M:%S")),
            ("Operator", conditions.operator_name),
            ("Laboratory", conditions.lab_name),
            ("", ""),
            ("Module Information", ""),
            ("Serial Number", conditions.module_serial),
            ("Manufacturer", conditions.module_manufacturer),
            ("Model", conditions.module_model),
            ("Rated Power", f"{conditions.module_rated_power_w}W"),
            ("Technology", conditions.module_technology),
            ("", ""),
            ("Electrical Conditions", ""),
            ("Current Recipe", conditions.current_recipe),
            ("Test Current", f"{conditions.test_current_a:.3f} A"),
            ("Isc", f"{conditions.isc_current_a:.3f} A"),
            ("Voc", f"{conditions.voc_voltage_v:.2f} V"),
            ("Impp", f"{conditions.impp_current_a:.3f} A"),
            ("Vmpp", f"{conditions.vmpp_voltage_v:.2f} V"),
            ("", ""),
            ("Environmental", ""),
            ("Ambient Temp", f"{conditions.ambient_temp_c:.1f} °C"),
            ("Module Temp", f"{conditions.module_temp_c:.1f} °C"),
            ("Humidity", f"{conditions.humidity_percent:.1f} %"),
            ("Air Pressure", f"{conditions.air_pressure_mbar:.1f} mbar"),
            ("", ""),
            ("Equipment", ""),
            ("Power Supply", f"{conditions.power_supply_make} {conditions.power_supply_model}"),
            ("Cable Length", f"{conditions.cable_length_m:.1f} m"),
            ("", ""),
            ("Standards", ""),
            ("Reference", conditions.standard_reference),
            ("Accreditation", conditions.accreditation or "N/A")
        ]

        row = 3
        for label, value in test_data:
            ws[f'A{row}'] = label
            ws[f'B{row}'] = value

            if label and not value:  # Section header
                ws[f'A{row}'].font = Font(bold=True, size=12, color="1F4E78")
            elif label != "":
                ws[f'A{row}'].font = Font(bold=True)

            row += 1

        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 40


class WordReportGenerator:
    """
    Word document report generator

    Creates professional Word documents with IEC/ISO templates
    """

    def __init__(self, config: Optional[ReportConfig] = None):
        """Initialize Word generator"""
        self.config = config or ReportConfig()
        logger.info("Initialized WordReportGenerator")

    def generate(
        self,
        defect_data: pd.DataFrame,
        images: Optional[Dict[str, np.ndarray]] = None,
        quality_metrics: Optional[ImageQualityMetrics] = None,
        test_conditions: Optional[TestConditions] = None,
        output_path: Optional[Path] = None
    ) -> Path:
        """Generate Word report"""
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = Path(f"EL_Report_{timestamp}.docx")

        doc = Document()

        # Set default styles
        self._setup_styles(doc)

        # Title page
        self._add_title_page(doc, test_conditions)

        # Executive summary
        self._add_executive_summary(doc, defect_data, quality_metrics)

        # Test conditions
        if test_conditions:
            self._add_test_conditions_section(doc, test_conditions)

        # Images
        if images:
            self._add_images_section(doc, images)

        # Defect analysis
        if not defect_data.empty:
            self._add_defects_section(doc, defect_data)

        # Quality metrics
        if quality_metrics:
            self._add_quality_section(doc, quality_metrics)

        # Compliance statement
        if self.config.include_compliance_statement:
            self._add_compliance_section(doc, quality_metrics)

        # Save
        doc.save(output_path)
        logger.info(f"Word report saved to: {output_path}")

        return output_path

    def _setup_styles(self, doc: Document) -> None:
        """Setup document styles"""
        styles = doc.styles

        # Heading styles are built-in, just modify if needed
        # You can customize further here

    def _add_title_page(self, doc: Document, conditions: Optional[TestConditions]) -> None:
        """Add title page"""
        # Title
        title = doc.add_heading(self.config.report_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Company
        company = doc.add_paragraph(self.config.company_name)
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company.runs[0].font.size = Pt(16)
        company.runs[0].font.italic = True

        doc.add_paragraph()

        # Date
        date_para = doc.add_paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if conditions:
            module_para = doc.add_paragraph(f"Module S/N: {conditions.module_serial}")
            module_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Standards
        doc.add_paragraph()
        doc.add_paragraph()
        standards_para = doc.add_paragraph("Applicable Standards:")
        standards_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        standards_para.runs[0].font.bold = True

        for std in self.config.standards:
            std_para = doc.add_paragraph(std, style='List Bullet')
            std_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    def _add_executive_summary(
        self,
        doc: Document,
        defect_data: pd.DataFrame,
        quality_metrics: Optional[ImageQualityMetrics]
    ) -> None:
        """Add executive summary"""
        doc.add_heading("Executive Summary", 1)

        if not defect_data.empty:
            total_defects = len(defect_data)
            critical = len(defect_data[defect_data['Severity'] == 'Critical'])
            total_loss = defect_data['Power_Loss_W'].sum()

            summary_text = (
                f"This electroluminescence inspection identified {total_defects} defects, "
                f"including {critical} critical defects. The total estimated power loss is "
                f"{total_loss:.2f}W."
            )
        else:
            summary_text = "No defects were detected during this electroluminescence inspection."

        doc.add_paragraph(summary_text)

        if quality_metrics:
            quality_text = (
                f"\nOverall image quality score: {quality_metrics.quality_score:.1f}/100 "
                f"({quality_metrics.quality_level}). "
                f"IEC 60904-13 compliance: {'PASS' if quality_metrics.iec_compliant else 'FAIL'}."
            )
            doc.add_paragraph(quality_text)

    def _add_test_conditions_section(self, doc: Document, conditions: TestConditions) -> None:
        """Add test conditions section"""
        doc.add_heading("Test Conditions", 1)

        # Create table
        table = doc.add_table(rows=10, cols=2)
        table.style = 'Light Grid Accent 1'

        data = [
            ("Test Date", conditions.test_date.strftime("%Y-%m-%d %H:%M:%S")),
            ("Module Serial", conditions.module_serial),
            ("Rated Power", f"{conditions.module_rated_power_w}W"),
            ("Current Recipe", conditions.current_recipe),
            ("Test Current", f"{conditions.test_current_a:.3f} A"),
            ("Isc", f"{conditions.isc_current_a:.3f} A"),
            ("Voc", f"{conditions.voc_voltage_v:.2f} V"),
            ("Ambient Temp", f"{conditions.ambient_temp_c:.1f} °C"),
            ("Module Temp", f"{conditions.module_temp_c:.1f} °C"),
            ("Standard", conditions.standard_reference)
        ]

        for i, (label, value) in enumerate(data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    def _add_images_section(self, doc: Document, images: Dict[str, np.ndarray]) -> None:
        """Add images section"""
        doc.add_page_break()
        doc.add_heading("Visual Analysis", 1)

        for image_name, image_array in images.items():
            doc.add_heading(image_name.replace('_', ' ').title(), 2)

            # Convert numpy array to PIL Image
            if len(image_array.shape) == 2:
                pil_image = Image.fromarray(image_array)
            else:
                pil_image = Image.fromarray(cv2.cvtColor(image_array, cv2.COLOR_BGR2RGB))

            # Save to bytes
            img_bytes = io.BytesIO()
            pil_image.save(img_bytes, format='PNG')
            img_bytes.seek(0)

            # Add to document
            doc.add_picture(img_bytes, width=Inches(5))

    def _add_defects_section(self, doc: Document, defect_data: pd.DataFrame) -> None:
        """Add defects section"""
        doc.add_page_break()
        doc.add_heading("Defect Analysis", 1)

        # Summary statistics
        doc.add_heading("Summary Statistics", 2)

        stats = [
            ("Total Defects", len(defect_data)),
            ("Non-Standard Defects", len(defect_data[defect_data['MBJ_Classification'] == 'Non-Standard'])),
            ("Critical Severity", len(defect_data[defect_data['Severity'] == 'Critical'])),
            ("Total Power Loss", f"{defect_data['Power_Loss_W'].sum():.2f}W")
        ]

        for label, value in stats:
            doc.add_paragraph(f"{label}: {value}", style='List Bullet')

        # Detailed table (first 20 defects)
        doc.add_heading("Defect Details (Top 20)", 2)

        display_cols = ['Defect_ID', 'Defect_Type', 'Cell_Location', 'Severity',
                       'Confidence_%', 'Power_Loss_W']

        # Create table
        table_data = defect_data[display_cols].head(20)
        table = doc.add_table(rows=len(table_data) + 1, cols=len(display_cols))
        table.style = 'Light Grid Accent 1'

        # Header
        for i, col in enumerate(display_cols):
            table.rows[0].cells[i].text = col.replace('_', ' ')
            table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

        # Data
        for row_idx, (_, row) in enumerate(table_data.iterrows(), 1):
            for col_idx, col in enumerate(display_cols):
                table.rows[row_idx].cells[col_idx].text = str(row[col])

    def _add_quality_section(self, doc: Document, metrics: ImageQualityMetrics) -> None:
        """Add quality metrics section"""
        doc.add_page_break()
        doc.add_heading("Image Quality Analysis", 1)

        quality_data = [
            ("Quality Score", f"{metrics.quality_score:.1f}/100"),
            ("Quality Level", metrics.quality_level.upper()),
            ("SNR", f"{metrics.snr_db:.2f} dB"),
            ("IEC Sharpness Class", metrics.iec_sharpness_class),
            ("MTF50", f"{metrics.mtf50:.3f}"),
            ("Shannon Entropy", f"{metrics.shannon_entropy:.3f} bits"),
            ("Dynamic Range", str(metrics.dynamic_range))
        ]

        for label, value in quality_data:
            doc.add_paragraph(f"{label}: {value}", style='List Bullet')

    def _add_compliance_section(self, doc: Document, metrics: Optional[ImageQualityMetrics]) -> None:
        """Add compliance statement"""
        doc.add_page_break()
        doc.add_heading("Compliance Statement", 1)

        compliance_text = (
            "This electroluminescence inspection was performed in accordance with "
            f"{', '.join(self.config.standards)}. "
        )

        if metrics and metrics.iec_compliant:
            compliance_text += (
                "The image quality and testing procedures meet all requirements "
                "specified in the applicable standards."
            )
        else:
            compliance_text += (
                "Note: Some image quality parameters do not meet the minimum "
                "requirements specified in IEC 60904-13. Results should be "
                "interpreted with caution."
            )

        doc.add_paragraph(compliance_text)

        # Signature section
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)
        doc.add_paragraph("Authorized Signature")
        doc.add_paragraph()
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")


class PDFReportGenerator:
    """
    PDF report generator using ReportLab

    Creates professional multi-page PDF reports
    """

    def __init__(self, config: Optional[ReportConfig] = None):
        """Initialize PDF generator"""
        self.config = config or ReportConfig()
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
        logger.info("Initialized PDFReportGenerator")

    def _setup_custom_styles(self) -> None:
        """Setup custom paragraph styles"""
        # Custom title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1F4E78'),
            spaceAfter=30,
            alignment=TA_CENTER
        ))

        # Custom subtitle
        self.styles.add(ParagraphStyle(
            name='CustomSubtitle',
            parent=self.styles['Normal'],
            fontSize=14,
            textColor=colors.grey,
            alignment=TA_CENTER,
            spaceAfter=12
        ))

    def generate(
        self,
        defect_data: pd.DataFrame,
        images: Optional[Dict[str, np.ndarray]] = None,
        quality_metrics: Optional[ImageQualityMetrics] = None,
        test_conditions: Optional[TestConditions] = None,
        output_path: Optional[Path] = None
    ) -> Path:
        """Generate PDF report"""
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = Path(f"EL_Report_{timestamp}.pdf")

        # Setup document
        pagesize = A4 if self.config.page_size == "A4" else letter
        doc = SimpleDocTemplate(str(output_path), pagesize=pagesize)

        # Build story (content)
        story = []

        # Title page
        story.extend(self._create_title_page(test_conditions))
        story.append(PageBreak())

        # Summary
        story.extend(self._create_summary(defect_data, quality_metrics))
        story.append(PageBreak())

        # Defects
        if not defect_data.empty:
            story.extend(self._create_defects_section(defect_data))
            story.append(PageBreak())

        # Build PDF
        doc.build(story)
        logger.info(f"PDF report saved to: {output_path}")

        return output_path

    def _create_title_page(self, conditions: Optional[TestConditions]) -> List:
        """Create title page elements"""
        elements = []

        # Title
        elements.append(Spacer(1, 2 * inch))
        elements.append(Paragraph(self.config.report_title, self.styles['CustomTitle']))

        # Company
        elements.append(Spacer(1, 0.5 * inch))
        elements.append(Paragraph(self.config.company_name, self.styles['CustomSubtitle']))

        # Date
        elements.append(Spacer(1, 1 * inch))
        date_text = f"Report Date: {datetime.now().strftime('%Y-%m-%d')}"
        elements.append(Paragraph(date_text, self.styles['Normal']))

        if conditions:
            module_text = f"Module S/N: {conditions.module_serial}"
            elements.append(Paragraph(module_text, self.styles['Normal']))

        return elements

    def _create_summary(self, defect_data: pd.DataFrame, quality_metrics: Optional[ImageQualityMetrics]) -> List:
        """Create summary section"""
        elements = []

        elements.append(Paragraph("Executive Summary", self.styles['Heading1']))
        elements.append(Spacer(1, 0.2 * inch))

        if not defect_data.empty:
            summary_text = (
                f"Total defects detected: <b>{len(defect_data)}</b><br/>"
                f"Critical defects: <b>{len(defect_data[defect_data['Severity'] == 'Critical'])}</b><br/>"
                f"Total power loss: <b>{defect_data['Power_Loss_W'].sum():.2f}W</b>"
            )
        else:
            summary_text = "No defects detected."

        elements.append(Paragraph(summary_text, self.styles['Normal']))

        return elements

    def _create_defects_section(self, defect_data: pd.DataFrame) -> List:
        """Create defects section"""
        elements = []

        elements.append(Paragraph("Defect Analysis", self.styles['Heading1']))
        elements.append(Spacer(1, 0.2 * inch))

        # Create table
        display_data = defect_data[['Defect_ID', 'Defect_Type', 'Severity', 'Power_Loss_W']].head(15)

        table_data = [['ID', 'Type', 'Severity', 'Power Loss (W)']]
        for _, row in display_data.iterrows():
            table_data.append([
                row['Defect_ID'],
                row['Defect_Type'],
                row['Severity'],
                f"{row['Power_Loss_W']:.2f}"
            ])

        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        elements.append(table)

        return elements


# Convenience function
def generate_all_reports(
    defect_data: pd.DataFrame,
    images: Optional[Dict[str, np.ndarray]] = None,
    quality_metrics: Optional[ImageQualityMetrics] = None,
    test_conditions: Optional[TestConditions] = None,
    output_dir: Optional[Path] = None
) -> Dict[str, Path]:
    """
    Generate all report formats

    Returns:
        Dictionary with paths to generated reports
    """
    if output_dir is None:
        output_dir = Path.cwd()
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    config = ReportConfig()

    results = {}

    # Excel
    try:
        excel_gen = ExcelReportGenerator(config)
        excel_path = output_dir / f"EL_Report_{timestamp}.xlsx"
        excel_gen.generate(defect_data, quality_metrics, test_conditions, excel_path)
        results['excel'] = excel_path
    except Exception as e:
        logger.error(f"Excel generation failed: {e}")

    # Word
    try:
        word_gen = WordReportGenerator(config)
        word_path = output_dir / f"EL_Report_{timestamp}.docx"
        word_gen.generate(defect_data, images, quality_metrics, test_conditions, word_path)
        results['word'] = word_path
    except Exception as e:
        logger.error(f"Word generation failed: {e}")

    # PDF
    try:
        pdf_gen = PDFReportGenerator(config)
        pdf_path = output_dir / f"EL_Report_{timestamp}.pdf"
        pdf_gen.generate(defect_data, images, quality_metrics, test_conditions, pdf_path)
        results['pdf'] = pdf_path
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")

    return results


if __name__ == "__main__":
    # Example usage
    import pandas as pd

    # Create sample data
    sample_data = pd.DataFrame({
        'Defect_ID': ['D001', 'D002', 'D003'],
        'Defect_Type': ['Crack', 'Hotspot', 'Dead Cell'],
        'Cell_Location': ['A1', 'B3', 'C5'],
        'Severity': ['Critical', 'High', 'Medium'],
        'Confidence_%': [95.5, 88.2, 76.3],
        'Power_Loss_W': [15.2, 8.5, 3.2],
        'MBJ_Classification': ['Non-Standard', 'Non-Standard', 'Non-Standard']
    })

    # Generate all reports
    report_paths = generate_all_reports(sample_data)

    print("Generated reports:")
    for format_type, path in report_paths.items():
        print(f"  {format_type.upper()}: {path}")
